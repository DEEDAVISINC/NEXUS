#!/usr/bin/env python3
"""
HIPAA Compliance Policy generator for Dee Davis Inc.

Reads company credentials from company_info.py.
Outputs .docx (Avenir where the font is available) and .pdf (Helvetica in ReportLab
if Avenir TTF is not bundled — Word output is the primary Avenir target).

Run: python3 hipaa_policy.py

Outputs: NEXUS DOCUMENTS/DDI-HIPAA-001_HIPAA_Compliance_Policy.{docx,pdf}
"""

from __future__ import annotations

import os
from datetime import date
from pathlib import Path

import company_info as ci

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as e:
    raise SystemExit("Install python-docx: pip install python-docx") from e

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )
except ImportError as e:
    raise SystemExit("Install reportlab: pip install reportlab") from e

# ─────────────────────────────────────────────────────────────────────────────
# Paths & constants
# ─────────────────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "NEXUS DOCUMENTS"
POLICY_NUMBER = "DDI-HIPAA-001"
EFFECTIVE_DATE = date.today()
FOOTER_LINE = (
    f"{ci.COMPANY_NAME} | {ci.ADDRESS_STREET}, {ci.ADDRESS_CITY}, {ci.ADDRESS_STATE} {ci.ADDRESS_ZIP} "
    f"| EIN: {ci.EIN} | Effective Date: {EFFECTIVE_DATE:%B %d, %Y} | Policy Number: {POLICY_NUMBER}"
)

# Prefer Avenir; Word substitutes if not installed on the machine.
DOC_FONT = "Avenir"
DOC_FONT_FALLBACK = "Calibri"


def _footer_text() -> str:
    return FOOTER_LINE


def policy_sections() -> list[tuple[str, list[str]]]:
    """Section title and body paragraphs (plain text)."""
    co = ci.COMPANY_NAME
    return [
        (
            "1. Purpose & Scope",
            [
                f"This HIPAA Compliance Policy establishes how {co} protects the privacy and security "
                "of Protected Health Information (PHI) in accordance with the Health Insurance Portability "
                "and Accountability Act of 1996 (HIPAA), the HIPAA Privacy Rule (45 CFR Part 160 and Subpart E of Part 164), "
                "the HIPAA Security Rule (45 CFR Part 160 and Subparts A and C of Part 164), "
                "the HIPAA Breach Notification Rule (45 CFR Part 164 Subpart D), and applicable state law where more stringent.",
                f"The policy applies to all workforce members (employees, contractors, volunteers, and others under {ci.COMPANY_NAME}'s direct control) "
                "who handle PHI in the course of work for or on behalf of the organization. It applies to PHI the organization "
                "creates, receives, maintains, or transmits in electronic, paper, or oral form.",
                f"Where {co} acts as a Business Associate, additional requirements in applicable Business Associate Agreements (BAAs) control "
                "if they are more specific than this policy; this policy must still be met at minimum.",
            ],
        ),
        (
            "2. Definitions",
            [
                "**Protected Health Information (PHI)** — Individually identifiable health information held or transmitted by a covered entity or business associate "
                "in any form or medium, except employment records held by a covered entity in its capacity as employer and certain other records excluded under HIPAA.",
                "**Business Associate Agreement (BAA)** — A written contract required by HIPAA (45 CFR 164.504(e)) that describes permitted uses and disclosures of PHI by a Business Associate, "
                "requires safeguards, and sets terms for breach notification and subcontractor compliance.",
                "**Covered Entity** — A health plan, health care clearinghouse, or health care provider who transmits any health information in electronic form "
                "in connection with a HIPAA-covered transaction (45 CFR 160.103).",
                "**Business Associate** — A person or entity that creates, receives, maintains, or transmits PHI on behalf of a covered entity (or another business associate) "
                "for a function or activity regulated by HIPAA, including claims processing, data analysis, utilization review, billing, legal, actuarial, consulting, "
                f"management, administrative, accreditation, or financial services (45 CFR 160.103). {co} may act as a Business Associate when performing such services for a Covered Entity.",
            ],
        ),
        (
            "3. Designated Privacy Officer",
            [
                f"{ci.OWNER_FULL_NAME} serves as the **Designated Privacy Officer** for {ci.COMPANY_NAME} and is accountable for HIPAA privacy compliance, "
                "including policy oversight, workforce training coordination, complaint handling, and liaison with regulatory authorities as appropriate.",
                f"**Title:** {ci.OWNER_TITLE}",
                f"**Organization:** {ci.COMPANY_NAME}",
                f"**Address:** {ci.ADDRESS_FULL}",
                f"**Phone:** {ci.PHONE_PRIMARY}",
                f"**Email:** {ci.EMAIL}",
                "The Privacy Officer may delegate operational tasks to qualified personnel but retains accountability for the HIPAA program.",
            ],
        ),
        (
            "4. Permitted Uses and Disclosures of PHI",
            [
                "PHI may be used or disclosed only as permitted or required by HIPAA and applicable BAAs:",
                "**Treatment, Payment, and Health Care Operations (TPO)** — As allowed under 45 CFR 164.506 when applicable to the relationship with the Covered Entity.",
                "**Authorization** — Uses and disclosures not otherwise permitted require a valid individual authorization unless an exception applies.",
                "**Required by law** — Disclosures required by statute, court order, or public health authority as permitted under 45 CFR 164.512.",
                "**Minimum Necessary** — All uses, disclosures, and requests must comply with Section 6 (Minimum Necessary Standard).",
                "Workforce members must not access, use, or disclose PHI for curiosity, personal reasons, or any purpose outside documented job duties.",
            ],
        ),
        (
            "5. Minimum Necessary Standard",
            [
                f"When using, disclosing, or requesting PHI, {co} must limit PHI to the **minimum necessary** to accomplish the intended purpose, except where the full record is required for treatment "
                "or as otherwise permitted by HIPAA (45 CFR 164.502(b), 164.514(d)).",
                "Supervisors and the Privacy Officer must ensure role-based access: workforce members receive access only to the PHI reasonably necessary to perform assigned duties.",
                "External requests for PHI must be reviewed to verify scope and legal basis before release.",
            ],
        ),
        (
            "6. Safeguards",
            [
                "**Administrative Safeguards** — Security management process; assigned security responsibility; workforce training; access management; contingency planning; "
                "evaluation; and Business Associate oversight, consistent with 45 CFR 164.308 where the Security Rule applies to electronic PHI.",
                "**Physical Safeguards** — Facility access controls; workstation use and security; device and media controls for PHI in physical form.",
                "**Technical Safeguards** — Access control; audit controls; integrity; person or entity authentication; and transmission security for electronic PHI (e-PHI) as applicable.",
                f"{co} maintains reasonable and appropriate safeguards based on size, complexity, and the nature of PHI handled.",
            ],
        ),
        (
            "7. Employee Training Requirements",
            [
                "All workforce members with access to PHI must receive **HIPAA privacy and security training** at hire and at least **annually** thereafter, "
                "or sooner when material changes to policies or regulations occur.",
                "Training must cover: this policy; PHI definitions; permitted uses; minimum necessary; safeguards; reporting incidents; and sanctions for violations.",
                "Documentation of training (attendance, date, topic) must be retained as required by the organization’s compliance program.",
            ],
        ),
        (
            "8. Business Associate Agreement Requirements",
            [
                f"Before a vendor or subcontractor may create, receive, maintain, or transmit PHI on behalf of {co} (as a Business Associate or subcontractor), "
                "a **written BAA** (or subcontractor agreement incorporating BAA terms) must be executed in accordance with 45 CFR 164.504(e).",
                f"The BAA must specify permitted uses and disclosures; require safeguards; require breach notification to {co}; and require flow-down obligations to subcontractors as required.",
                "No PHI may be shared with a new Business Associate until the BAA is fully executed, except as allowed by law in limited circumstances.",
            ],
        ),
        (
            "9. Breach Notification Procedures (60-Day Rule and Related Requirements)",
            [
                "A **breach** is the acquisition, access, use, or disclosure of PHI in a manner not permitted under the Privacy Rule that compromises security or privacy, "
                "subject to the risk assessment and exceptions in 45 CFR 164.402–164.406.",
                "**Discovery and internal escalation** — Any workforce member who suspects a breach must notify the **Privacy Officer immediately** (same-day if feasible).",
                "**Risk assessment** — The Privacy Officer (or designee) conducts a documented risk assessment to determine if unsecured PHI was compromised and whether notification is required.",
                f"**Notification to the Covered Entity** — If {co} is a Business Associate, notification to the relevant Covered Entity must be **without unreasonable delay** and "
                "no later than **60 calendar days** after discovery of the breach (45 CFR 164.410), unless the BAA specifies an earlier timeframe — the stricter timeline applies.",
                f"**HHS and individuals** — Covered Entity clients are responsible for notifying affected individuals and, when required, the Secretary of HHS and media; "
                f"{co} will cooperate fully and provide information necessary for those notifications per the BAA.",
                "**Documentation** — Breach documentation (assessment, facts, mitigation) must be retained for at least six years as required by HIPAA.",
            ],
        ),
        (
            "10. Member / Individual Rights",
            [
                f"To the extent {co} maintains PHI on behalf of a Covered Entity, **individual rights** (access, amendment, accounting of disclosures, restrictions, "
                f"confidential communications) are primarily exercised through the Covered Entity; {co} will **forward requests promptly** to the Covered Entity "
                "and assist as required by the BAA.",
                f"If {co} receives a request directly, the Privacy Officer will acknowledge receipt, coordinate with the Covered Entity within required timeframes, "
                "and document the response path.",
            ],
        ),
        (
            "11. Complaint Procedures",
            [
                f"Individuals or workforce members may file privacy complaints **in writing** to the Privacy Officer at {ci.EMAIL} or {ci.ADDRESS_FULL}.",
                "Complaints will be acknowledged; investigated reasonably; and documented. **No retaliation** is permitted against good-faith reporting.",
                "Individuals may also file complaints with the **U.S. Department of Health and Human Services, Office for Civil Rights (OCR)**. "
                "Contact information for OCR is available at www.hhs.gov/ocr/privacy/hipaa/complaints/.",
            ],
        ),
        (
            "12. Sanctions for Non-Compliance",
            [
                "Violations of this policy or HIPAA may result in **disciplinary action** up to and including termination of employment or contract, "
                "and referral to law enforcement where appropriate.",
                "Intentional misuse of PHI, snooping, or failure to report suspected breaches may result in immediate escalation.",
                f"{ci.COMPANY_NAME} may also be subject to **civil and criminal penalties** under HIPAA and related law for non-compliance.",
            ],
        ),
        (
            "13. Policy Review",
            [
                "This policy will be reviewed **at least annually** and updated when laws, regulations, or business operations change materially.",
            ],
        ),
    ]


def _format_paragraph_plain(text: str) -> str:
    """Strip markdown-style bold markers for plain Word/PDF text."""
    return text.replace("**", "")


def build_docx(path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        fp.text = ""
        fr = fp.add_run(_footer_text())
        fr.font.size = Pt(8)
        fr.font.name = DOC_FONT
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Styles
    normal = doc.styles["Normal"]
    normal.font.name = DOC_FONT
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("HIPAA COMPLIANCE POLICY")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.name = DOC_FONT

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run(f"{ci.COMPANY_NAME}\n")
    sr.font.size = Pt(12)
    sr.font.name = DOC_FONT
    sr2 = st.add_run(
        f"Policy Number: {POLICY_NUMBER}\nEffective Date: {EFFECTIVE_DATE:%B %d, %Y}"
    )
    sr2.font.size = Pt(11)
    sr2.font.name = DOC_FONT

    doc.add_paragraph()

    for sec_title, paras in policy_sections():
        h = doc.add_heading(sec_title, level=1)
        for r in h.runs:
            r.font.name = DOC_FONT
            r.font.size = Pt(12)
        for ptext in paras:
            plain = _format_paragraph_plain(ptext)
            para = doc.add_paragraph(plain)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.name = DOC_FONT
                run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("ACKNOWLEDGED AND ADOPTED:\n\n").bold = True
    doc.add_paragraph()
    line = doc.add_paragraph()
    line.add_run("_" * 50)
    name = doc.add_paragraph()
    nr = name.add_run(f"{ci.OWNER_FULL_NAME}\n{ci.OWNER_TITLE}\n{ci.COMPANY_NAME}")
    nr.font.name = DOC_FONT
    for para in (sig, line, name):
        for r in para.runs:
            r.font.name = DOC_FONT

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_pdf(path: Path) -> None:
    """PDF using ReportLab; Helvetica for reliability (Avenir via TTF optional)."""
    path.parent.mkdir(parents=True, exist_ok=True)

    def draw_footer(canvas, doc_template):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#333333"))
        text = _footer_text()
        canvas.drawCentredString(LETTER[0] / 2, 0.45 * inch, text)
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name="PolicyTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        alignment=1,
        spaceAfter=12,
    )
    sub_style = ParagraphStyle(
        name="PolicySub",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        alignment=1,
        spaceAfter=18,
    )
    h1_style = ParagraphStyle(
        name="H1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=12,
        spaceBefore=14,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        name="Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )
    sig_style = ParagraphStyle(
        name="Sig",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        spaceBefore=24,
    )

    story = []
    story.append(Paragraph("HIPAA COMPLIANCE POLICY", title_style))
    story.append(
        Paragraph(
            f"{ci.COMPANY_NAME}<br/>Policy Number: {POLICY_NUMBER}<br/>"
            f"Effective Date: {EFFECTIVE_DATE:%B %d, %Y}",
            sub_style,
        )
    )
    story.append(Spacer(1, 0.15 * inch))

    for sec_title, paras in policy_sections():
        story.append(Paragraph(sec_title.replace("&", "&amp;"), h1_style))
        for ptext in paras:
            plain = _format_paragraph_plain(ptext)
            # Escape for ReportLab XML-ish parser
            esc = (
                plain.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            story.append(Paragraph(esc, body_style))

    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph("<b>ACKNOWLEDGED AND ADOPTED:</b>", sig_style))
    story.append(Spacer(1, 0.4 * inch))
    story.append(Paragraph("_" * 50, sig_style))
    story.append(
        Paragraph(
            f"{ci.OWNER_FULL_NAME}<br/>{ci.OWNER_TITLE}<br/>{ci.COMPANY_NAME}",
            sig_style,
        )
    )

    doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    stem = f"{POLICY_NUMBER}_HIPAA_Compliance_Policy"
    docx_path = OUTPUT_DIR / f"{stem}.docx"
    pdf_path = OUTPUT_DIR / f"{stem}.pdf"
    print(f"Writing {docx_path} ...")
    build_docx(docx_path)
    print(f"Writing {pdf_path} ...")
    build_pdf(pdf_path)
    print("Done.")


if __name__ == "__main__":
    main()
