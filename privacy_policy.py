#!/usr/bin/env python3
"""
Privacy Policy generator for Dee Davis Inc.

Reads credentials from company_info.py.
Outputs .docx (Avenir where available) and .pdf (Helvetica via ReportLab).

Run: python3 privacy_policy.py

Outputs: NEXUS DOCUMENTS/DDI-PRIV-001_Privacy_Policy.{docx,pdf}
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

import company_info as ci

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as e:
    raise SystemExit("Install python-docx: pip install python-docx") from e

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
except ImportError as e:
    raise SystemExit("Install reportlab: pip install reportlab") from e

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "NEXUS DOCUMENTS"
POLICY_NUMBER = "DDI-PRIV-001"
EFFECTIVE_DATE = date.today()

# Footer per specification (every page)
FOOTER_LINE = (
    f"{ci.COMPANY_NAME} | {POLICY_NUMBER} | {ci.ADDRESS_STREET}, "
    f"{ci.ADDRESS_CITY}, {ci.ADDRESS_STATE} {ci.ADDRESS_ZIP}"
)

PRIVACY_EMAIL = ci.EMAIL  # info@deedavis.biz per company_info / user request

DOC_FONT = "Avenir"


def _footer_text() -> str:
    return FOOTER_LINE


def _format_paragraph_plain(text: str) -> str:
    return text.replace("**", "")


def policy_sections() -> list[tuple[str, list[str]]]:
    co = ci.COMPANY_NAME
    return [
        (
            "1. Introduction",
            [
                f"This Privacy Policy describes how **{co}** (“we,” “us,” or “our”) collects, uses, discloses, and protects information "
                "in connection with our websites, services, and business operations. We are committed to protecting privacy and "
                "complying with applicable laws, including the Health Insurance Portability and Accountability Act (HIPAA) and "
                "state privacy requirements, where they apply to the information we handle.",
                "By using our services or website, you acknowledge this policy. If you do not agree, please do not use our services. "
                "We may update this policy as described in Section 10.",
            ],
        ),
        (
            "2. What Information We Collect",
            [
                "**Contact and identity information** — Name, title, organization, mailing address, email address, telephone number, "
                "and similar identifiers you or your organization provide when requesting services, registering, or communicating with us.",
                "**Protected Health Information (PHI)** — When we provide or coordinate health-related services as permitted by law or contract, "
                "we may create, receive, maintain, or transmit PHI (for example, information relating to past, present, or future physical or mental health "
                "or payment for care) only as necessary to perform agreed services and in accordance with HIPAA and applicable Business Associate Agreements.",
                "**Payment and billing information** — Billing address, payment card or bank details, invoice references, and transaction records needed to "
                "process payments and fulfill contracts. Payment card processing may be handled by third-party processors; we do not store full card data "
                "except as needed for reconciliation and as permitted by law.",
                "**Operational and technical data** — Device identifiers, IP address, browser type, pages viewed, and similar data when you use our website "
                "(see Section 7).",
                "**Other information** — Information you voluntarily provide in forms, surveys, email, or phone calls relevant to our services.",
            ],
        ),
        (
            "3. How We Use Information",
            [
                "We use collected information to:",
                "**Provide services** — Deliver contract management, logistics, healthcare-adjacent coordination, and related operations you or your organization request.",
                "**Communicate** — Respond to inquiries, send service-related notices, and manage our relationship with clients, partners, and vendors.",
                "**Billing and collections** — Process payments, invoices, and account administration.",
                "**Compliance and safety** — Meet legal and regulatory obligations; protect rights, privacy, safety, and security; detect and prevent fraud or misuse.",
                "**Improvement** — Analyze usage in aggregate to improve our website and operations (where not prohibited by law or contract).",
                "We do not sell personal information for monetary consideration. Uses of PHI follow HIPAA, our policies, and applicable BAAs.",
            ],
        ),
        (
            "4. Who We Share Information With",
            [
                "We disclose information only as needed to operate our business and as permitted by law:",
                "**Managed Care Organizations (MCOs) and Medicaid programs** — When we contract or coordinate with health plans, state Medicaid agencies, "
                "or related entities, we may share information necessary to perform services, as directed by the client or required by program rules, "
                "subject to HIPAA and program-specific agreements.",
                "**Subcontractors and service providers** — Vetted partners who perform work on our behalf (for example, transportation, laboratory, "
                "technology, or professional services) under written agreements that require appropriate safeguards and, where applicable, HIPAA-compliant "
                "Business Associate or subcontractor terms.",
                "**Government and legal** — When required by law, regulation, court order, or lawful governmental request.",
                "**Professional advisors** — Attorneys, accountants, or insurers under confidentiality obligations when necessary.",
                "We require third parties to use information only for the purposes we authorize and to implement reasonable protections.",
            ],
        ),
        (
            "5. How We Protect Information",
            [
                f"{co} maintains **administrative, physical, and technical safeguards** appropriate to the nature of the information and our operations, including:",
                "Access controls and workforce training; secure handling of physical records; and technical measures for systems and electronic communications "
                "where we control them.",
                "No method of transmission or storage is completely secure; we strive to use commercially reasonable measures and to comply with HIPAA Security Rule "
                "requirements where we handle electronic PHI as a Business Associate.",
                "Breaches involving PHI are handled in accordance with HIPAA breach notification rules and our internal procedures.",
            ],
        ),
        (
            "6. Member and Client Rights",
            [
                "Depending on your relationship with us and applicable law, you may have rights regarding your information, including:",
                "**Access** — Request a copy of or access to certain personal or health information we maintain about you, as required by HIPAA or state law.",
                "**Correction** — Request amendment or correction of inaccurate information, subject to legal limits and our role (for example, when we hold PHI "
                "on behalf of a Covered Entity, requests may be routed through that entity as required by HIPAA).",
                "**Deletion** — Where applicable law provides a right to delete personal information and we are not required to retain it for legal or contractual reasons, "
                "we will process requests in accordance with law. HIPAA and healthcare retention rules may limit deletion of certain records.",
                "**Complaints** — You may file a complaint with us (see Section 9) or, for HIPAA-related concerns, with the U.S. Department of Health and Human Services Office for Civil Rights.",
                "To exercise rights, contact the Privacy Officer using the information below. We will respond within timeframes required by applicable law.",
            ],
        ),
        (
            "7. Cookies and Website Data",
            [
                f"Our website at **{ci.WEBSITE}** may use cookies, local storage, or similar technologies to remember preferences, maintain sessions, "
                "and understand aggregate traffic patterns.",
                "You may control cookies through your browser settings; disabling cookies may limit certain site features.",
                "We do not use cookies to knowingly collect PHI from casual website browsing unless you submit it through a secure form or portal designed for that purpose.",
            ],
        ),
        (
            "8. Third-Party Services",
            [
                "Our website or operations may link to or integrate with third-party sites, analytics tools, payment processors, cloud hosting, or communication platforms. "
                "Those services have their own privacy policies; we are not responsible for their practices except as required by law or contract.",
                "We encourage you to read third-party policies before providing information to them.",
            ],
        ),
        (
            "9. Privacy Officer — How to Contact Us",
            [
                f"**Privacy Officer:** {ci.OWNER_FULL_NAME}",
                f"**Email:** {PRIVACY_EMAIL}",
                f"**Phone:** {ci.PHONE_PRIMARY}",
                f"**Mailing address:** {ci.ADDRESS_FULL}",
                "For privacy requests, questions, or complaints, please email the Privacy Officer at the address above or write to our mailing address. "
                "Include sufficient detail for us to verify and respond to your request.",
            ],
        ),
        (
            "10. Effective Date and Policy Updates",
            [
                f"**Effective date:** {EFFECTIVE_DATE:%B %d, %Y}.",
                f"We may revise this Privacy Policy from time to time. The current version will be posted with an updated effective date. "
                f"Material changes may be communicated through our website or direct notice where appropriate. Continued use of our services after the effective date "
                "constitutes acceptance of the updated policy where permitted by law.",
                f"**Policy number:** {POLICY_NUMBER}.",
            ],
        ),
    ]


def build_docx(path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        fp.text = ""
        fr = fp.add_run(_footer_text())
        fr.font.size = Pt(8)
        fr.font.name = DOC_FONT
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    normal = doc.styles["Normal"]
    normal.font.name = DOC_FONT
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("PRIVACY POLICY")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.name = DOC_FONT

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run(f"{ci.COMPANY_NAME}\n")
    sr.font.size = Pt(12)
    sr.font.name = DOC_FONT
    sr2 = st.add_run(
        f"Policy Number: {POLICY_NUMBER}\nEffective Date: {EFFECTIVE_DATE:%B %d, %Y}"
    )
    sr2.font.size = Pt(11)
    sr2.font.name = DOC_FONT

    doc.add_paragraph()

    for sec_title, paras in policy_sections():
        h = doc.add_heading(sec_title, level=1)
        for r in h.runs:
            r.font.name = DOC_FONT
            r.font.size = Pt(12)
        for ptext in paras:
            plain = _format_paragraph_plain(ptext)
            para = doc.add_paragraph(plain)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.name = DOC_FONT
                run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("APPROVED BY:\n\n").bold = True
    doc.add_paragraph()
    line = doc.add_paragraph()
    line.add_run("_" * 50)
    name = doc.add_paragraph()
    nr = name.add_run(f"{ci.OWNER_FULL_NAME}\n{ci.OWNER_TITLE}\n{ci.COMPANY_NAME}")
    nr.font.name = DOC_FONT
    for para in (sig, line, name):
        for r in para.runs:
            r.font.name = DOC_FONT

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)

    def draw_footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#333333"))
        canvas.drawCentredString(LETTER[0] / 2, 0.45 * inch, _footer_text())
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name="PrivTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        alignment=1,
        spaceAfter=12,
    )
    sub_style = ParagraphStyle(
        name="PrivSub",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        alignment=1,
        spaceAfter=18,
    )
    h1_style = ParagraphStyle(
        name="PrivH1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=12,
        spaceBefore=14,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        name="PrivBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )
    sig_style = ParagraphStyle(
        name="PrivSig",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        spaceBefore=24,
    )

    story = [
        Paragraph("PRIVACY POLICY", title_style),
        Paragraph(
            f"{ci.COMPANY_NAME}<br/>Policy Number: {POLICY_NUMBER}<br/>"
            f"Effective Date: {EFFECTIVE_DATE:%B %d, %Y}",
            sub_style,
        ),
        Spacer(1, 0.15 * inch),
    ]

    for sec_title, paras in policy_sections():
        story.append(Paragraph(sec_title.replace("&", "&amp;"), h1_style))
        for ptext in paras:
            plain = _format_paragraph_plain(ptext)
            esc = plain.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(esc, body_style))

    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph("<b>APPROVED BY:</b>", sig_style))
    story.append(Spacer(1, 0.4 * inch))
    story.append(Paragraph("_" * 50, sig_style))
    story.append(
        Paragraph(
            f"{ci.OWNER_FULL_NAME}<br/>{ci.OWNER_TITLE}<br/>{ci.COMPANY_NAME}",
            sig_style,
        )
    )

    doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    stem = f"{POLICY_NUMBER}_Privacy_Policy"
    docx_path = OUTPUT_DIR / f"{stem}.docx"
    pdf_path = OUTPUT_DIR / f"{stem}.pdf"
    print(f"Writing {docx_path} ...")
    build_docx(docx_path)
    print(f"Writing {pdf_path} ...")
    build_pdf(pdf_path)
    print("Done.")


if __name__ == "__main__":
    main()
