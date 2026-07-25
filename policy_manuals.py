#!/usr/bin/env python3
"""
Generate PDF/DOCX copies of GATEWAY workforce policy manuals for MCO audit packs.

Source of truth for workforce reading/e-sign remains:
  gateway-portal/policies/*.html on gateway.deedavis.biz

Run: python3 policy_manuals.py

Outputs under NEXUS DOCUMENTS/:
  DDI-HR-001_Employee_Handbook.{docx,pdf}
  DDI-HR-002_Contractor_Obligations_Handbook.{docx,pdf}
  DDI-FDR-001_FDR_Compliance_Policies_Procedures.{docx,pdf}
  DDI-HR-003_Code_of_Conduct_Conflict_of_Interest.{docx,pdf}
  DDI-HR-004_Confidentiality_NDA_Acknowledgment.{docx,pdf}
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

import company_info as ci

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as e:
    raise SystemExit("Install python-docx: pip install python-docx") from e

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
except ImportError as e:
    raise SystemExit("Install reportlab: pip install reportlab") from e

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "NEXUS DOCUMENTS"
EFFECTIVE_DATE = date(2026, 7, 25)
DOC_FONT = "Avenir"
VERSION = "1.0"

POLICIES = [
    {
        "id": "DDI-HR-001",
        "title": "EMPLOYEE HANDBOOK",
        "stem": "DDI-HR-001_Employee_Handbook",
        "portal": "gateway.deedavis.biz/policies/employee-handbook",
        "sections": [
            ("1. Purpose & Scope", [
                f"{ci.COMPANY_NAME} (“DDI”) is a nationwide contract management Third-Party Administrator (TPA) "
                f"headquartered at {ci.ADDRESS_FULL}. This Employee Handbook sets expectations for W-2 employees.",
                "Employees must also complete any division-required policies (including FDR Compliance Policies & Procedures where applicable), HIPAA training, and GATEWAY onboarding steps.",
            ]),
            ("2. Employment Relationship", [
                "Employment with Dee Davis Inc. is at-will to the extent permitted by applicable law. This handbook is not a contract of employment and does not guarantee continued employment.",
            ]),
            ("3. Equal Opportunity & Anti-Harassment", [
                "DDI provides equal employment opportunity and prohibits unlawful discrimination and harassment. Report concerns to your manager, HR, or gc@deedavis.biz. Retaliation for good-faith reports is prohibited.",
            ]),
            ("4. Conduct & Professional Standards", [
                "Represent DDI professionally. Follow the “coordinate, not provide” TPA standard. Never disclose end-buyer, agency, or solicitation identity to suppliers unless expressly authorized. Protect system credentials. Escalate compliance and member-risk issues immediately.",
            ]),
            ("5. Confidentiality, PHI & PII", [
                "Protect confidential business information and any PHI/PII. See DDI-HIPAA-001 and DDI-PRIV-001. Unauthorized disclosure may result in discipline up to termination.",
            ]),
            ("6. Timekeeping, Pay & Expenses", [
                "Record time accurately. Falsifying time or expenses is grounds for discipline. Company email aliases remain Company property.",
            ]),
            ("7. Systems Access & GATEWAY", [
                "GATEWAY (gateway.deedavis.biz) is DDI’s workforce onboarding and acknowledgment portal. Required training, uploads, and policy acknowledgments are tracked in GATEWAY and NEXUS.",
            ]),
            ("8. Compliance Training & Screening", [
                "Complete assigned compliance training within stated deadlines. Where the assigned division requires FDR screening, OIG LEIE and GSA SAM.gov checks apply per DDI-FDR-001. Failure to remain current may block work via the can-work gate.",
            ]),
            ("9. Conflicts of Interest", [
                "Disclose actual or potential conflicts under DDI-HR-003 and re-attest as required.",
            ]),
            ("10. Discipline & Separation", [
                "DDI may take corrective action up to termination for policy or compliance failures. Upon separation, return Company property and cease use of systems and confidential information.",
            ]),
            ("11. Acknowledgment", [
                "Typed-name electronic signature in GATEWAY for DDI-HR-001 confirms the employee has read and agrees to follow this handbook (Version 1.0).",
            ]),
        ],
    },
    {
        "id": "DDI-HR-002",
        "title": "CONTRACTOR OBLIGATIONS HANDBOOK",
        "stem": "DDI-HR-002_Contractor_Obligations_Handbook",
        "portal": "gateway.deedavis.biz/policies/contractor-handbook",
        "sections": [
            ("1. Purpose & Scope", [
                "This handbook describes contractual flow-down obligations for independent contractors engaged by Dee Davis Inc. It does not create an employment relationship.",
            ]),
            ("2. Independent Contractor Status", [
                "Contractors are not W-2 employees. DDI does not run E-Verify/I-9 on contractors. Payment is deliverable/milestone-based per the written agreement.",
            ]),
            ("3. Scope Discipline", [
                "Perform only work within the written scope. Do not expand into PHI or member-facing tasks unless expressly included and required training/BAA conditions are met.",
            ]),
            ("4. Confidentiality & Buyer Protection", [
                "Protect DDI confidential information and authorized PHI/PII. Never reveal end-buyer or solicitation identity to unauthorized parties. See DDI-HR-004 and applicable BAA.",
            ]),
            ("5. Conduct", [
                "Follow DDI-HR-003 as a contractual flow-down. Report suspected FWA through channels in DDI-FDR-001 when the engagement is FDR-adjacent.",
            ]),
            ("6. Systems Access", [
                "Access is minimum-necessary and engagement-specific. GATEWAY is used for document upload, training, and policy e-sign. Access ends when the engagement ends.",
            ]),
            ("7. Training & Screening", [
                "Complete training marked applicable to the engagement. FDR-adjacent divisions require exclusion screening and DDI-FDR-001 acknowledgment.",
            ]),
            ("8. Insurance & Classification", [
                "Provide Certificate of Insurance when required. Cooperate with worker-classification documentation.",
            ]),
            ("9. Acknowledgment", [
                "Typed-name electronic signature in GATEWAY for DDI-HR-002 confirms agreement to these flow-down obligations (Version 1.0).",
            ]),
        ],
    },
    {
        "id": "DDI-FDR-001",
        "title": "FDR COMPLIANCE POLICIES & PROCEDURES",
        "stem": "DDI-FDR-001_FDR_Compliance_Policies_Procedures",
        "portal": "gateway.deedavis.biz/policies/fdr-compliance",
        "sections": [
            ("1. Purpose & Scope", [
                "Where Dee Davis Inc. is a First Tier, Downstream, or Related Entity (FDR) under Medicare/Medicaid managed care contracts, these Policies & Procedures apply to workforce in FDR-adjacent divisions.",
                "Division-driven applicability: DEPOINTE (NEMT Coordination), HAVEN, SHIELD, VITAL, and Corporate/HR/Admin require OIG LEIE + GSA SAM screening and this acknowledgment. Freight 1st Direct, 3D Ink Signatures/CNTDA, DEPOINTE DNA, ARENA/PRIME, and unassigned do not by default.",
            ]),
            ("2. Compliance Responsibility & Reporting", [
                f"Ultimate responsibility rests with the {ci.OWNER_TITLE}. Report FWA or compliance concerns to gc@deedavis.biz; quality/ops issues to qc@deedavis.biz. Good-faith reporters are protected from retaliation.",
            ]),
            ("3. Compliance & FWA Training", [
                "Assigned CMS FDR curriculum is tracked in GATEWAY. DDI internal target: 30 days. CMS hard floor for General Compliance/FWA and Medicare Fraud & Abuse: 90 days. Missing the 90-day floor fails the can-work gate. Recurrence follows the GATEWAY training catalog.",
            ]),
            ("4. Exclusion Screening", [
                "For FDR-applicable divisions: OIG LEIE and GSA SAM.gov screening at hire/engagement and monthly thereafter. Open flagged matches and overdue/never-screened status block can-work until resolved.",
            ]),
            ("5. Record Retention", [
                "FDR-related training, screening, acknowledgment, and audit records are retained for a minimum of ten (10) years. GATEWAY records are archived, not hard-deleted.",
            ]),
            ("6. Annual FDR Compliance Attestation", [
                "DDI maintains an organization-level Annual FDR Compliance Attestation on a calendar-year cycle, tracked in NEXUS HR FDR ATTESTATION.",
            ]),
            ("7. Downstream Entities", [
                "DDI remains responsible for contract management and compliance oversight when using subcontractors or fulfillment partners on FDR-adjacent work.",
            ]),
            ("8. Related Policies", [
                "DDI-HIPAA-001, DDI-PRIV-001, DDI-HR-001/002, DDI-HR-003, DDI-HR-004.",
            ]),
            ("9. Discipline", [
                "Non-compliance may result in removal from FDR-adjacent work, employment discipline, contractor termination, and required reporting.",
            ]),
            ("10. Acknowledgment", [
                "Typed-name electronic signature in GATEWAY for DDI-FDR-001 confirms the signer has read and agrees to comply (Version 1.0).",
            ]),
        ],
    },
    {
        "id": "DDI-HR-003",
        "title": "CODE OF CONDUCT / CONFLICT OF INTEREST POLICY",
        "stem": "DDI-HR-003_Code_of_Conduct_Conflict_of_Interest",
        "portal": "gateway.deedavis.biz/policies/code-of-conduct",
        "sections": [
            ("1. Purpose", [
                "Defines ethical conduct and conflict-of-interest disclosure for employees and contractors (contractual flow-down for contractors).",
            ]),
            ("2. Standards of Conduct", [
                "Comply with law, contracts, and DDI policies. No harassment or unlawful discrimination. No FWA. Protect confidential information and PHI/PII. No improper gifts or kickbacks. No misrepresentation. Never reveal end-buyer identity to unauthorized parties.",
            ]),
            ("3. Conflicts of Interest", [
                "Disclose actual or potential conflicts in writing to your manager/Engagement Manager and gc@deedavis.biz before the conflict affects work. Recusal or mitigation may be required.",
            ]),
            ("4. Reporting & Non-Retaliation", [
                "Report violations in good faith without fear of retaliation.",
            ]),
            ("5. Violations", [
                "May result in employment discipline, contractor termination, removal from FDR-adjacent work, and reporting as required.",
            ]),
            ("6. Acknowledgment", [
                "Typed-name electronic signature in GATEWAY for DDI-HR-003 confirms agreement (Version 1.0).",
            ]),
        ],
    },
    {
        "id": "DDI-HR-004",
        "title": "CONFIDENTIALITY / NDA ACKNOWLEDGMENT",
        "stem": "DDI-HR-004_Confidentiality_NDA_Acknowledgment",
        "portal": "gateway.deedavis.biz/policies/nda",
        "sections": [
            ("1. Confidential Information", [
                "Non-public business, technical, financial, pricing, supplier, client, member, solicitation, proposal, and operational information of Dee Davis Inc. and its clients.",
            ]),
            ("2. Obligations", [
                "Use Confidential Information only for authorized DDI duties. Do not disclose to unauthorized persons. Protect credentials and report suspected breaches to gc@deedavis.biz. Return or securely destroy Confidential Information upon request or separation, subject to legal retention duties.",
            ]),
            ("3. PHI / HIPAA", [
                "If the role involves PHI, follow DDI-HIPAA-001 and any applicable Business Associate Agreement.",
            ]),
            ("4. Duration", [
                "Obligations continue during employment/engagement and survive afterward while information remains non-public, or longer if a separate NDA or law requires.",
            ]),
            ("5. Remedies", [
                "Unauthorized disclosure may cause irreparable harm; DDI may seek injunctive and other remedies.",
            ]),
            ("6. Acknowledgment (E-Sign)", [
                "Typed name, timestamp, and IP in GATEWAY constitute the electronic signature for DDI-HR-004 (Version 1.0). Separately executed NDAs/BAAs control on conflicting terms.",
            ]),
        ],
    },
]


def _footer(policy_id: str, portal: str) -> str:
    return (
        f"{ci.COMPANY_NAME} | {ci.ADDRESS_FULL} | "
        f"Effective: {EFFECTIVE_DATE:%B %d, %Y} | {policy_id} v{VERSION} | {portal}"
    )


def build_docx(policy: dict, path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        fp.text = ""
        fr = fp.add_run(_footer(policy["id"], policy["portal"]))
        fr.font.size = Pt(7)
        fr.font.name = DOC_FONT
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    normal = doc.styles["Normal"]
    normal.font.name = DOC_FONT
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(policy["title"])
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.name = DOC_FONT

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run(
        f"{ci.COMPANY_NAME}\nPolicy Number: {policy['id']}\n"
        f"Version: {VERSION}\nEffective Date: {EFFECTIVE_DATE:%B %d, %Y}\n"
        f"Controlled copy: {policy['portal']}"
    )
    sr.font.size = Pt(10)
    sr.font.name = DOC_FONT
    doc.add_paragraph()

    for sec_title, paras in policy["sections"]:
        h = doc.add_heading(sec_title, level=1)
        for r in h.runs:
            r.font.name = DOC_FONT
            r.font.size = Pt(12)
        for ptext in paras:
            para = doc.add_paragraph(ptext)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.name = DOC_FONT
                run.font.size = Pt(11)

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("ADOPTED FOR COMPANY USE:\n\n").bold = True
    line = doc.add_paragraph()
    line.add_run("_" * 50)
    name = doc.add_paragraph()
    name.add_run(f"{ci.OWNER_FULL_NAME}\n{ci.OWNER_TITLE}\n{ci.COMPANY_NAME}")
    for para in (sig, line, name):
        for r in para.runs:
            r.font.name = DOC_FONT

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_pdf(policy: dict, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    footer_text = _footer(policy["id"], policy["portal"])

    def draw_footer(canvas, doc_template):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#333333"))
        canvas.drawCentredString(LETTER[0] / 2, 0.4 * inch, footer_text[:120])
        if len(footer_text) > 120:
            canvas.drawCentredString(LETTER[0] / 2, 0.28 * inch, footer_text[120:240])
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=0.8 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name="PolicyTitle", parent=styles["Heading1"],
        fontName="Helvetica-Bold", fontSize=14, alignment=1, spaceAfter=10,
    )
    sub_style = ParagraphStyle(
        name="PolicySub", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, alignment=1, spaceAfter=14,
    )
    h1_style = ParagraphStyle(
        name="H1", parent=styles["Heading1"],
        fontName="Helvetica-Bold", fontSize=11, spaceBefore=12, spaceAfter=6,
    )
    body_style = ParagraphStyle(
        name="Body", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, leading=13,
        alignment=TA_JUSTIFY, spaceAfter=7,
    )
    sig_style = ParagraphStyle(
        name="Sig", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, spaceBefore=18,
    )

    def esc(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = []
    story.append(Paragraph(esc(policy["title"]), title_style))
    story.append(Paragraph(
        f"{esc(ci.COMPANY_NAME)}<br/>Policy Number: {esc(policy['id'])}<br/>"
        f"Version: {VERSION}<br/>Effective Date: {EFFECTIVE_DATE:%B %d, %Y}<br/>"
        f"Controlled copy: {esc(policy['portal'])}",
        sub_style,
    ))
    for sec_title, paras in policy["sections"]:
        story.append(Paragraph(esc(sec_title), h1_style))
        for ptext in paras:
            story.append(Paragraph(esc(ptext), body_style))
    story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph("<b>ADOPTED FOR COMPANY USE:</b>", sig_style))
    story.append(Spacer(1, 0.35 * inch))
    story.append(Paragraph("_" * 50, sig_style))
    story.append(Paragraph(
        f"{esc(ci.OWNER_FULL_NAME)}<br/>{esc(ci.OWNER_TITLE)}<br/>{esc(ci.COMPANY_NAME)}",
        sig_style,
    ))
    doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for policy in POLICIES:
        docx_path = OUTPUT_DIR / f"{policy['stem']}.docx"
        pdf_path = OUTPUT_DIR / f"{policy['stem']}.pdf"
        print(f"Writing {docx_path.name} ...")
        build_docx(policy, docx_path)
        print(f"Writing {pdf_path.name} ...")
        build_pdf(policy, pdf_path)
    print("Done — all GATEWAY policy manuals exported.")


if __name__ == "__main__":
    main()
